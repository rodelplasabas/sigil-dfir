"""
SIGIL Report Generator — Automated DOCX report generation from DFIR findings.
Produces professional Compromise Assessment / VAPT-style reports.
"""

import os
import io
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn


# ── Color Palette ─────────────────────────────────────────────────────────────
COLORS = {
    "critical": RGBColor(0xEF, 0x44, 0x44),
    "high": RGBColor(0xF9, 0x73, 0x16),
    "medium": RGBColor(0xF5, 0x9E, 0x0B),
    "low": RGBColor(0x3B, 0x82, 0xF6),
    "info": RGBColor(0x6B, 0x72, 0x80),
    "clean": RGBColor(0x10, 0xB9, 0x81),
    "suspicious": RGBColor(0xF5, 0x9E, 0x0B),
    "compromised": RGBColor(0xEF, 0x44, 0x44),
    "dark": RGBColor(0x0F, 0x17, 0x2A),
    "header_bg": RGBColor(0x1E, 0x29, 0x3B),
    "accent": RGBColor(0x38, 0xBD, 0xF8),
    "text": RGBColor(0x33, 0x33, 0x33),
    "muted": RGBColor(0x6B, 0x72, 0x80),
    "white": RGBColor(0xFF, 0xFF, 0xFF),
}

SEVERITY_ORDER = {"critical": 0, "high": 1, "medium": 2, "low": 3, "info": 4}


def _set_cell_shading(cell, color_hex: str):
    """Set cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color_hex,
        qn("w:val"): "clear",
    })
    shading.append(shading_elm)


def _add_styled_run(paragraph, text: str, bold=False, italic=False, size=None, color=None, font_name=None):
    """Add a styled text run to a paragraph."""
    run = paragraph.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if font_name:
        run.font.name = font_name
    return run


def _add_heading(doc, text: str, level: int = 1):
    """Add a heading with consistent styling."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = COLORS["dark"]
    return heading


def _add_severity_badge(paragraph, severity: str):
    """Add a colored severity badge inline."""
    color = COLORS.get(severity, COLORS["muted"])
    _add_styled_run(paragraph, f"  [{severity.upper()}]", bold=True, size=10, color=color)


def _build_findings_table(doc, findings: list):
    """Build the findings summary table."""
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    headers = ["#", "Finding", "Severity", "Confidence", "MITRE ATT&CK"]
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        _add_styled_run(p, header, bold=True, size=9, color=COLORS["white"], font_name="Arial")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_shading(cell, "1E293B")

    # Data rows
    for idx, f in enumerate(findings, 1):
        row = table.add_row()
        severity = f.get("severity", "medium")
        sev_color = COLORS.get(severity, COLORS["muted"])

        # #
        cell0 = row.cells[0]
        cell0.text = ""
        p0 = cell0.paragraphs[0]
        _add_styled_run(p0, str(idx), size=9, font_name="Arial")
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Finding name + ID
        cell1 = row.cells[1]
        cell1.text = ""
        p1 = cell1.paragraphs[0]
        _add_styled_run(p1, f.get("name", "Unknown"), bold=True, size=9, font_name="Arial")
        _add_styled_run(p1, f"  ({f.get('id', '')})", size=8, color=COLORS["muted"], font_name="Arial")

        # Severity
        cell2 = row.cells[2]
        cell2.text = ""
        p2 = cell2.paragraphs[0]
        _add_styled_run(p2, severity.upper(), bold=True, size=9, color=sev_color, font_name="Arial")
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Confidence
        cell3 = row.cells[3]
        cell3.text = ""
        p3 = cell3.paragraphs[0]
        _add_styled_run(p3, f"{f.get('confidence', 0)}%", size=9, font_name="Arial")
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # MITRE
        cell4 = row.cells[4]
        cell4.text = ""
        p4 = cell4.paragraphs[0]
        mitre = f.get("mitre", [])
        _add_styled_run(p4, ", ".join(mitre) if mitre else "N/A", size=8, font_name="Consolas")

        # Alternate row shading
        if idx % 2 == 0:
            for cell in row.cells:
                _set_cell_shading(cell, "F8FAFC")

    # Set column widths
    widths = [Cm(1), Cm(6), Cm(2.2), Cm(2.2), Cm(4)]
    for row in table.rows:
        for i, width in enumerate(widths):
            row.cells[i].width = width

    return table


def generate_report(
    case_meta: dict,
    findings: list,
    overall_score: dict,
    artifacts: list,
    ioc_list: list = None,
) -> io.BytesIO:
    """
    Generate a DOCX report from SIGIL findings.

    Returns BytesIO buffer containing the DOCX file.
    """
    doc = Document()

    # ── Page setup ────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # ── Default font ──────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style.font.color.rgb = COLORS["text"]
    style.paragraph_format.space_after = Pt(6)

    # ── Title Page ────────────────────────────────────────────────────────
    # Spacer
    for _ in range(4):
        doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_styled_run(title, "SIGIL", bold=True, size=36, color=COLORS["accent"], font_name="Arial")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_styled_run(subtitle, "DFIR Compromise Assessment Report", bold=True, size=18, color=COLORS["dark"], font_name="Arial")

    doc.add_paragraph("")

    # Overall score badge
    score_label = overall_score.get("label", "CLEAN") if overall_score else "CLEAN"
    score_color = {
        "CLEAN": COLORS["clean"],
        "SUSPICIOUS": COLORS["suspicious"],
        "COMPROMISED": COLORS["compromised"]
    }.get(score_label, COLORS["muted"])

    score_p = doc.add_paragraph()
    score_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_styled_run(score_p, f"Overall Assessment: ", size=14, color=COLORS["text"], font_name="Arial")
    _add_styled_run(score_p, score_label, bold=True, size=16, color=score_color, font_name="Arial")

    doc.add_paragraph("")

    # Case metadata table
    case_name = case_meta.get("name", "Untitled Case") if case_meta else "Untitled Case"
    examiner = case_meta.get("examiner", "N/A") if case_meta else "N/A"
    description = case_meta.get("description", "") if case_meta else ""
    report_date = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    report_scope = case_meta.get("report_scope", "Full Report") if case_meta else "Full Report"

    meta_items = [
        ("Case Name", case_name),
        ("Examiner", examiner),
        ("Report Date", report_date),
        ("Report Scope", report_scope),
        ("Artifacts Analyzed", str(len(artifacts)) if artifacts else "0"),
    ]

    meta_table = doc.add_table(rows=len(meta_items), cols=2)
    meta_table.style = "Table Grid"
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (label, value) in enumerate(meta_items):
        cell_label = meta_table.rows[i].cells[0]
        cell_label.text = ""
        _add_styled_run(cell_label.paragraphs[0], label, bold=True, size=10, font_name="Arial")
        _set_cell_shading(cell_label, "F1F5F9")
        cell_label.width = Cm(4)

        cell_value = meta_table.rows[i].cells[1]
        cell_value.text = ""
        _add_styled_run(cell_value.paragraphs[0], value, size=10, font_name="Arial")
        cell_value.width = Cm(11)

    if description:
        doc.add_paragraph("")
        desc_p = doc.add_paragraph()
        desc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_styled_run(desc_p, description, italic=True, size=10, color=COLORS["muted"])

    # Page break
    doc.add_page_break()

    # ── Executive Summary ─────────────────────────────────────────────────
    _add_heading(doc, "1. Executive Summary", level=1)

    sorted_findings = sorted(findings, key=lambda f: SEVERITY_ORDER.get(f.get("severity", "info"), 4))
    crit_count = sum(1 for f in findings if f.get("severity") == "critical")
    high_count = sum(1 for f in findings if f.get("severity") == "high")
    med_count = sum(1 for f in findings if f.get("severity") == "medium")
    low_count = sum(1 for f in findings if f.get("severity") == "low")
    total_events = sum(len(f.get("matched_events", [])) for f in findings)

    summary_p = doc.add_paragraph()
    _add_styled_run(summary_p,
        f"This report presents the findings of a DFIR compromise assessment conducted using SIGIL. "
        f"A total of {len(artifacts) if artifacts else 0} artifact(s) were analyzed, yielding "
        f"{len(findings)} detection finding(s) across {total_events} matched events. "
        f"The overall assessment is ",
        size=11)
    _add_styled_run(summary_p, score_label, bold=True, size=11, color=score_color)
    _add_styled_run(summary_p, ".", size=11)

    # Severity breakdown
    if findings:
        breakdown_p = doc.add_paragraph()
        _add_styled_run(breakdown_p, "Severity Breakdown: ", bold=True, size=11)
        parts = []
        if crit_count: parts.append(f"{crit_count} Critical")
        if high_count: parts.append(f"{high_count} High")
        if med_count: parts.append(f"{med_count} Medium")
        if low_count: parts.append(f"{low_count} Low")
        _add_styled_run(breakdown_p, ", ".join(parts) + ".", size=11)

    # IOC summary
    if ioc_list and len(ioc_list) > 0:
        ioc_ips = [i for i in ioc_list if i.get("type") == "ip"]
        ioc_domains = [i for i in ioc_list if i.get("type") == "domain"]
        ioc_findings = [f for f in findings if f.get("is_ioc_rule")]
        ioc_p = doc.add_paragraph()
        _add_styled_run(ioc_p,
            f"IOC Hunting: {len(ioc_list)} indicators were loaded ({len(ioc_ips)} IPs, {len(ioc_domains)} domains). "
            f"{len(ioc_findings)} IOC-related finding(s) were detected.",
            size=11)

    # ── Artifacts Analyzed ────────────────────────────────────────────────
    _add_heading(doc, "2. Artifacts Analyzed", level=1)

    if artifacts:
        art_table = doc.add_table(rows=1, cols=4)
        art_table.style = "Table Grid"
        art_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, header in enumerate(["#", "Filename", "Log Type", "Events"]):
            cell = art_table.rows[0].cells[i]
            cell.text = ""
            _add_styled_run(cell.paragraphs[0], header, bold=True, size=9, color=COLORS["white"], font_name="Arial")
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_cell_shading(cell, "1E293B")

        for idx, art in enumerate(artifacts, 1):
            row = art_table.add_row()
            row.cells[0].text = str(idx)
            row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row.cells[1].text = art.get("name", "Unknown")
            log_type_display = {
                "windows_event_log": "Windows Event Log",
                "web_server_log": "Web Server Log",
                "registry": "Registry"
            }.get(art.get("log_type", ""), art.get("log_type", "Unknown"))
            row.cells[2].text = log_type_display
            row.cells[3].text = str(art.get("event_count", art.get("eventCount", "N/A")))
            row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)
                        run.font.name = "Arial"

            if idx % 2 == 0:
                for cell in row.cells:
                    _set_cell_shading(cell, "F8FAFC")

        # File hashes section
        has_hashes = any(art.get("hashes") for art in artifacts)
        if has_hashes:
            doc.add_paragraph("")
            hash_heading = doc.add_heading("File Hashes", level=2)
            for run in hash_heading.runs:
                run.font.color.rgb = COLORS["dark"]

            hash_table = doc.add_table(rows=1, cols=5)
            hash_table.style = "Table Grid"
            hash_table.alignment = WD_TABLE_ALIGNMENT.CENTER

            for i, header in enumerate(["Filename", "Size", "MD5", "SHA1", "SHA256"]):
                cell = hash_table.rows[0].cells[i]
                cell.text = ""
                _add_styled_run(cell.paragraphs[0], header, bold=True, size=8, color=COLORS["white"], font_name="Arial")
                _set_cell_shading(cell, "1E293B")

            for art in artifacts:
                hashes = art.get("hashes")
                if not hashes:
                    continue
                row = hash_table.add_row()

                # Filename
                row.cells[0].text = ""
                _add_styled_run(row.cells[0].paragraphs[0], art.get("name", ""), size=7, font_name="Arial")

                # Size
                file_size = hashes.get("file_size", 0)
                if file_size > 1048576:
                    size_str = f"{file_size / 1048576:.1f} MB"
                elif file_size > 1024:
                    size_str = f"{file_size / 1024:.1f} KB"
                else:
                    size_str = f"{file_size} B"
                row.cells[1].text = ""
                _add_styled_run(row.cells[1].paragraphs[0], size_str, size=7, font_name="Consolas")

                # Hashes
                for col, algo in [(2, "md5"), (3, "sha1"), (4, "sha256")]:
                    row.cells[col].text = ""
                    _add_styled_run(row.cells[col].paragraphs[0], hashes.get(algo, "N/A"), size=6, font_name="Consolas")

    else:
        doc.add_paragraph("No artifacts were provided for analysis.")

    # ── Findings Summary ──────────────────────────────────────────────────
    _add_heading(doc, "3. Findings Summary", level=1)

    if findings:
        _build_findings_table(doc, sorted_findings)
    else:
        doc.add_paragraph("No findings were detected during the analysis.")

    # ── Detailed Findings ─────────────────────────────────────────────────
    _add_heading(doc, "4. Detailed Findings", level=1)

    for idx, f in enumerate(sorted_findings, 1):
        severity = f.get("severity", "medium")
        sev_color = COLORS.get(severity, COLORS["muted"])

        # Finding header
        finding_title = doc.add_heading(level=2)
        _add_styled_run(finding_title, f"4.{idx}  ", size=13, color=COLORS["dark"])
        _add_styled_run(finding_title, f.get("name", "Unknown Finding"), bold=True, size=13, color=COLORS["dark"])
        _add_styled_run(finding_title, f"  [{severity.upper()}]", bold=True, size=11, color=sev_color)

        # Rule ID + confidence
        meta_p = doc.add_paragraph()
        _add_styled_run(meta_p, f"Rule ID: ", bold=True, size=10, color=COLORS["muted"])
        _add_styled_run(meta_p, f.get("id", "N/A"), size=10, font_name="Consolas", color=COLORS["accent"])
        _add_styled_run(meta_p, f"    Confidence: ", bold=True, size=10, color=COLORS["muted"])
        _add_styled_run(meta_p, f"{f.get('confidence', 0)}%", size=10, font_name="Consolas")
        _add_styled_run(meta_p, f"    Matched Events: ", bold=True, size=10, color=COLORS["muted"])
        _add_styled_run(meta_p, str(len(f.get("matched_events", []))), size=10, font_name="Consolas")

        # MITRE tags
        mitre = f.get("mitre", [])
        if mitre:
            mitre_p = doc.add_paragraph()
            _add_styled_run(mitre_p, "MITRE ATT&CK: ", bold=True, size=10, color=COLORS["muted"])
            _add_styled_run(mitre_p, ", ".join(mitre), size=10, font_name="Consolas", color=COLORS["accent"])

        # Description
        desc = f.get("description", "")
        if desc:
            doc.add_paragraph(desc)

        # Recommendations / Next Steps
        next_steps = f.get("next_steps", [])
        if next_steps:
            rec_p = doc.add_paragraph()
            _add_styled_run(rec_p, "Recommendations:", bold=True, size=11)
            for step in next_steps:
                step_p = doc.add_paragraph(style="List Bullet")
                _add_styled_run(step_p, str(step), size=10)

        # Evidence sample (first 5 matched events)
        matched = f.get("matched_events", [])
        if matched:
            ev_heading = doc.add_paragraph()
            _add_styled_run(ev_heading, f"Evidence Sample ({min(len(matched), 5)} of {len(matched)} events):", bold=True, size=10)

            ev_table = doc.add_table(rows=1, cols=4)
            ev_table.style = "Table Grid"
            ev_table.alignment = WD_TABLE_ALIGNMENT.CENTER

            for i, header in enumerate(["Record ID", "Event ID", "Timestamp", "Content (excerpt)"]):
                cell = ev_table.rows[0].cells[i]
                cell.text = ""
                _add_styled_run(cell.paragraphs[0], header, bold=True, size=8, color=COLORS["white"], font_name="Arial")
                _set_cell_shading(cell, "374151")

            for ev in matched[:5]:
                row = ev_table.add_row()
                row.cells[0].text = str(ev.get("record_id", ""))
                row.cells[1].text = str(ev.get("event_id", ""))
                row.cells[2].text = str(ev.get("timestamp", ""))[:19]
                content = str(ev.get("content", ev.get("message", "")))[:150]
                row.cells[3].text = content

                for cell in row.cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.size = Pt(7)
                            run.font.name = "Consolas"

        doc.add_paragraph("")  # Spacer between findings

    # ── IOC List ──────────────────────────────────────────────────────────
    if ioc_list and len(ioc_list) > 0:
        _add_heading(doc, "5. Indicators of Compromise (IOCs)", level=1)

        ioc_table = doc.add_table(rows=1, cols=3)
        ioc_table.style = "Table Grid"
        ioc_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, header in enumerate(["#", "Type", "Value"]):
            cell = ioc_table.rows[0].cells[i]
            cell.text = ""
            _add_styled_run(cell.paragraphs[0], header, bold=True, size=9, color=COLORS["white"], font_name="Arial")
            _set_cell_shading(cell, "1E293B")

        for idx, ioc in enumerate(ioc_list, 1):
            row = ioc_table.add_row()
            row.cells[0].text = str(idx)
            row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row.cells[1].text = ioc.get("type", "unknown").upper()
            row.cells[2].text = ioc.get("value", "")
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)
                        run.font.name = "Consolas"

    # ── Footer ────────────────────────────────────────────────────────────
    doc.add_page_break()
    footer_heading = doc.add_paragraph()
    footer_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_styled_run(footer_heading, "— End of Report —", italic=True, size=10, color=COLORS["muted"])

    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_styled_run(footer_p, f"Generated by SIGIL DFIR Compromise Assessment Tool on {report_date}", size=9, color=COLORS["muted"])

    # ── Save to buffer ────────────────────────────────────────────────────
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer